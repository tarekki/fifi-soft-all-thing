"""
Reports Views
عروض التقارير

هذا الملف يحتوي على جميع views للتقارير.
This file contains all Reports views.

Endpoints:
- GET /api/v1/admin/reports/sales/       - Sales report
- GET /api/v1/admin/reports/products/    - Products report
- GET /api/v1/admin/reports/users/        - Users report
- GET /api/v1/admin/reports/commissions/  - Commissions report
- GET /api/v1/admin/reports/export/       - Export report as Word
"""

from rest_framework import status
from rest_framework.views import APIView
from rest_framework.response import Response
from django.utils import timezone
from django.utils.translation import gettext_lazy as _
from django.db.models import Sum, Count, Avg, Q, F
from django.db.models.functions import TruncDate
from datetime import timedelta, datetime
from decimal import Decimal
from drf_spectacular.utils import extend_schema, OpenApiParameter, OpenApiResponse
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from admin_api.permissions import IsAdminUser
from admin_api.throttling import AdminUserRateThrottle
from admin_api.serializers.reports import (
    SalesReportSerializer,
    ProductsReportSerializer,
    UsersReportSerializer,
    CommissionsReportSerializer,
)
from core.utils import success_response, error_response


# =============================================================================
# Helper Functions
# دوال مساعدة
# =============================================================================

def get_date_range(date_range: str):
    """
    Get date range from period string.
    الحصول على نطاق التاريخ من سلسلة الفترة.
    
    Args:
        date_range: '7days', '30days', '90days', 'year', or 'custom'
    
    Returns:
        tuple: (date_from, date_to)
    """
    now = timezone.now()
    today = now.date()
    
    if date_range == '7days':
        date_from = today - timedelta(days=7)
        date_to = today
    elif date_range == '30days':
        date_from = today - timedelta(days=30)
        date_to = today
    elif date_range == '90days':
        date_from = today - timedelta(days=90)
        date_to = today
    elif date_range == 'year':
        date_from = today.replace(month=1, day=1)
        date_to = today
    else:  # custom or default to 30 days
        date_from = today - timedelta(days=30)
        date_to = today
    
    return date_from, date_to


def get_previous_period(date_from, date_to):
    """
    Get previous period for comparison.
    الحصول على الفترة السابقة للمقارنة.
    """
    period_days = (date_to - date_from).days
    prev_date_to = date_from - timedelta(days=1)
    prev_date_from = prev_date_to - timedelta(days=period_days)
    return prev_date_from, prev_date_to


def calculate_change(current, previous):
    """
    Calculate percentage change.
    حساب نسبة التغيير.
    """
    if previous == 0:
        return 100.0 if current > 0 else 0.0
    return float(((current - previous) / previous) * 100)


# =============================================================================
# Sales Report View
# عرض تقرير المبيعات
# =============================================================================

class SalesReportView(APIView):
    """
    Get sales report.
    الحصول على تقرير المبيعات.
    """
    
    permission_classes = [IsAdminUser]
    throttle_classes = [AdminUserRateThrottle]
    
    @extend_schema(
        summary='Sales Report',
        description='Get sales report with statistics and daily sales data',
        parameters=[
            OpenApiParameter(
                name='date_range',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Date range: 7days, 30days, 90days, year',
                enum=['7days', '30days', '90days', 'year'],
                default='30days'
            ),
            OpenApiParameter(
                name='date_from',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Start date (YYYY-MM-DD) for custom range',
            ),
            OpenApiParameter(
                name='date_to',
                type=str,
                location=OpenApiParameter.QUERY,
                description='End date (YYYY-MM-DD) for custom range',
            ),
        ],
        responses={
            200: SalesReportSerializer,
        },
        tags=['Admin Reports'],
    )
    def get(self, request):
        """
        Get sales report.
        الحصول على تقرير المبيعات.
        """
        from orders.models import Order
        from django.contrib.auth import get_user_model
        
        User = get_user_model()
        
        # Get date range
        date_range = request.query_params.get('date_range', '30days')
        
        if date_range == 'custom':
            date_from_str = request.query_params.get('date_from')
            date_to_str = request.query_params.get('date_to')
            if date_from_str and date_to_str:
                date_from = datetime.strptime(date_from_str, '%Y-%m-%d').date()
                date_to = datetime.strptime(date_to_str, '%Y-%m-%d').date()
            else:
                date_from, date_to = get_date_range('30days')
        else:
            date_from, date_to = get_date_range(date_range)
        
        # Convert to datetime for filtering
        date_from_dt = timezone.make_aware(datetime.combine(date_from, datetime.min.time()))
        date_to_dt = timezone.make_aware(datetime.combine(date_to, datetime.max.time()))
        
        # Get previous period for comparison
        prev_date_from, prev_date_to = get_previous_period(date_from, date_to)
        prev_date_from_dt = timezone.make_aware(datetime.combine(prev_date_from, datetime.min.time()))
        prev_date_to_dt = timezone.make_aware(datetime.combine(prev_date_to, datetime.max.time()))
        
        # Current period orders
        current_orders = Order.objects.filter(
            created_at__gte=date_from_dt,
            created_at__lte=date_to_dt,
            status__in=['delivered', 'confirmed', 'shipped']
        )
        
        # Previous period orders
        previous_orders = Order.objects.filter(
            created_at__gte=prev_date_from_dt,
            created_at__lte=prev_date_to_dt,
            status__in=['delivered', 'confirmed', 'shipped']
        )
        
        # Calculate statistics
        total_revenue = current_orders.aggregate(total=Sum('total'))['total'] or Decimal('0.00')
        total_orders = current_orders.count()
        avg_order_value = (total_revenue / total_orders) if total_orders > 0 else Decimal('0.00')
        
        # Previous period statistics
        prev_revenue = previous_orders.aggregate(total=Sum('total'))['total'] or Decimal('0.00')
        prev_orders = previous_orders.count()
        prev_avg_order_value = (prev_revenue / prev_orders) if prev_orders > 0 else Decimal('0.00')
        
        # New users in period
        new_users = User.objects.filter(
            is_staff=False,
            date_joined__gte=date_from_dt,
            date_joined__lte=date_to_dt
        ).count()
        
        prev_new_users = User.objects.filter(
            is_staff=False,
            date_joined__gte=prev_date_from_dt,
            date_joined__lte=prev_date_to_dt
        ).count()
        
        # Daily sales
        daily_sales_data = current_orders.annotate(
            date=TruncDate('created_at')
        ).values('date').annotate(
            sales=Sum('total')
        ).order_by('date')
        
        # Map to day names (Arabic)
        day_names = {
            0: 'الاثنين',
            1: 'الثلاثاء',
            2: 'الأربعاء',
            3: 'الخميس',
            4: 'الجمعة',
            5: 'السبت',
            6: 'الأحد',
        }
        
        daily_sales = []
        for item in daily_sales_data:
            day_name = day_names.get(item['date'].weekday(), item['date'].strftime('%Y-%m-%d'))
            daily_sales.append({
                'day': day_name,
                'sales': item['sales'] or Decimal('0.00')
            })
        
        # Get detailed orders list
        detailed_orders = current_orders.select_related('user').prefetch_related('items')[:100]  # Limit to 100 orders
        
        orders_list = []
        status_display_map = {
            'pending': _('قيد الانتظار / Pending'),
            'confirmed': _('مؤكد / Confirmed'),
            'shipped': _('تم الشحن / Shipped'),
            'delivered': _('تم التوصيل / Delivered'),
            'cancelled': _('ملغي / Cancelled'),
        }
        
        for order in detailed_orders:
            orders_list.append({
                'id': order.id,
                'order_number': order.order_number or f"ORD-{order.id:06d}",
                'customer_name': order.customer_name or (order.user.get_full_name() if order.user else _('ضيف / Guest')),
                'customer_phone': getattr(order, 'customer_phone', ''),
                'status': order.status,
                'status_display': status_display_map.get(order.status, order.status),
                'total': order.total,
                'items_count': order.items.count() if hasattr(order, 'items') else 0,
                'created_at': order.created_at,
            })
        
        # Build response
        data = {
            'total_revenue': total_revenue,
            'total_orders': total_orders,
            'avg_order_value': avg_order_value,
            'new_users': new_users,
            'revenue_change': calculate_change(float(total_revenue), float(prev_revenue)),
            'orders_change': calculate_change(total_orders, prev_orders),
            'avg_order_value_change': calculate_change(float(avg_order_value), float(prev_avg_order_value)),
            'new_users_change': calculate_change(new_users, prev_new_users),
            'daily_sales': daily_sales,
            'orders': orders_list,
            'date_from': date_from,
            'date_to': date_to,
        }
        
        # Serialize the data
        serializer = SalesReportSerializer(data=data)
        if serializer.is_valid():
            return success_response(
                data=serializer.validated_data,
                message=_('تم جلب تقرير المبيعات / Sales report retrieved')
            )
        else:
            # If validation fails, return data directly (shouldn't happen with calculated data)
            return success_response(
                data=data,
                message=_('تم جلب تقرير المبيعات / Sales report retrieved')
            )


# =============================================================================
# Products Report View
# عرض تقرير المنتجات
# =============================================================================

class ProductsReportView(APIView):
    """
    Get products report.
    الحصول على تقرير المنتجات.
    """
    
    permission_classes = [IsAdminUser]
    throttle_classes = [AdminUserRateThrottle]
    
    @extend_schema(
        summary='Products Report',
        description='Get products report with top products and sales by category',
        parameters=[
            OpenApiParameter(
                name='date_range',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Date range: 7days, 30days, 90days, year',
                enum=['7days', '30days', '90days', 'year'],
                default='30days'
            ),
        ],
        responses={
            200: ProductsReportSerializer,
        },
        tags=['Admin Reports'],
    )
    def get(self, request):
        """
        Get products report.
        الحصول على تقرير المنتجات.
        """
        from orders.models import Order, OrderItem
        from products.models import Product, Category
        
        date_range = request.query_params.get('date_range', '30days')
        date_from, date_to = get_date_range(date_range)
        date_from_dt = timezone.make_aware(datetime.combine(date_from, datetime.min.time()))
        date_to_dt = timezone.make_aware(datetime.combine(date_to, datetime.max.time()))
        
        # Get orders in period
        orders = Order.objects.filter(
            created_at__gte=date_from_dt,
            created_at__lte=date_to_dt,
            status__in=['delivered', 'confirmed', 'shipped']
        )
        
        # Top products by sales with details
        order_items = OrderItem.objects.filter(
            order__in=orders
        ).select_related(
            'product_variant__product__vendor',
            'product_variant__product__category'
        ).values(
            'product_variant__product__id',
            'product_variant__product__name',
            'product_variant__product__vendor__name',
            'product_variant__product__category__name_ar',
            'product_variant__product__category__name'
        ).annotate(
            sales=Count('id'),
            revenue=Sum(F('price') * F('quantity'))
        ).order_by('-sales')[:20]  # Top 20 products
        
        # Get stock quantities for products
        from products.models import ProductVariant
        product_ids = [item['product_variant__product__id'] for item in order_items]
        stock_data = ProductVariant.objects.filter(
            product_id__in=product_ids
        ).values('product_id').annotate(
            total_stock=Sum('stock_quantity')
        )
        stock_map = {item['product_id']: item['total_stock'] for item in stock_data}
        
        top_products = []
        for item in order_items:
            product_id = item['product_variant__product__id']
            category_name = item['product_variant__product__category__name_ar'] or item['product_variant__product__category__name'] or 'أخرى'
            top_products.append({
                'id': product_id,
                'name': item['product_variant__product__name'],
                'vendor_name': item['product_variant__product__vendor__name'] or 'غير معروف',
                'category_name': category_name,
                'sales': item['sales'],
                'revenue': item['revenue'] or Decimal('0.00'),
                'stock_quantity': stock_map.get(product_id, 0)
            })
        
        # Sales by category
        category_sales = OrderItem.objects.filter(
            order__in=orders
        ).values(
            'product_variant__product__category__name_ar',
            'product_variant__product__category__name'
        ).annotate(
            sales=Sum(F('price') * F('quantity'))
        ).order_by('-sales')
        
        total_category_sales = sum(item['sales'] or Decimal('0.00') for item in category_sales)
        
        sales_by_category = []
        for item in category_sales:
            category_name = item['product_variant__product__category__name_ar'] or item['product_variant__product__category__name'] or 'أخرى'
            sales_amount = item['sales'] or Decimal('0.00')
            percentage = float((sales_amount / total_category_sales) * 100) if total_category_sales > 0 else 0.0
            
            sales_by_category.append({
                'category': category_name,
                'sales': sales_amount,
                'percentage': round(percentage, 1)
            })
        
        data = {
            'top_products': top_products,
            'sales_by_category': sales_by_category,
            'date_from': date_from,
            'date_to': date_to,
        }
        
        # Serialize the data
        serializer = ProductsReportSerializer(data=data)
        if serializer.is_valid():
            return success_response(
                data=serializer.validated_data,
                message=_('تم جلب تقرير المنتجات / Products report retrieved')
            )
        else:
            # If validation fails, return data directly
            return success_response(
                data=data,
                message=_('تم جلب تقرير المنتجات / Products report retrieved')
            )


# =============================================================================
# Users Report View
# عرض تقرير المستخدمين
# =============================================================================

class UsersReportView(APIView):
    """
    Get users report.
    الحصول على تقرير المستخدمين.
    """
    
    permission_classes = [IsAdminUser]
    throttle_classes = [AdminUserRateThrottle]
    
    @extend_schema(
        summary='Users Report',
        description='Get users report with statistics',
        parameters=[
            OpenApiParameter(
                name='date_range',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Date range: 7days, 30days, 90days, year',
                enum=['7days', '30days', '90days', 'year'],
                default='30days'
            ),
        ],
        responses={
            200: UsersReportSerializer,
        },
        tags=['Admin Reports'],
    )
    def get(self, request):
        """
        Get users report.
        الحصول على تقرير المستخدمين.
        """
        from django.contrib.auth import get_user_model
        
        User = get_user_model()
        
        date_range = request.query_params.get('date_range', '30days')
        date_from, date_to = get_date_range(date_range)
        date_from_dt = timezone.make_aware(datetime.combine(date_from, datetime.min.time()))
        date_to_dt = timezone.make_aware(datetime.combine(date_to, datetime.max.time()))
        
        # Get previous period
        prev_date_from, prev_date_to = get_previous_period(date_from, date_to)
        prev_date_from_dt = timezone.make_aware(datetime.combine(prev_date_from, datetime.min.time()))
        prev_date_to_dt = timezone.make_aware(datetime.combine(prev_date_to, datetime.max.time()))
        
        # New users in period
        new_users = User.objects.filter(
            is_staff=False,
            date_joined__gte=date_from_dt,
            date_joined__lte=date_to_dt
        ).count()
        
        prev_new_users = User.objects.filter(
            is_staff=False,
            date_joined__gte=prev_date_from_dt,
            date_joined__lte=prev_date_to_dt
        ).count()
        
        # Total and active users
        total_users = User.objects.filter(is_staff=False).count()
        active_users = User.objects.filter(
            is_staff=False,
            is_active=True
        ).count()
        
        # Get detailed users list with their orders
        from orders.models import Order
        detailed_users = User.objects.filter(
            is_staff=False,
            date_joined__gte=date_from_dt,
            date_joined__lte=date_to_dt
        ).annotate(
            orders_count=Count('orders'),
            total_spent=Sum('orders__total', filter=Q(orders__status__in=['delivered', 'confirmed', 'shipped']))
        )[:100]  # Limit to 100 users
        
        users_list = []
        for user in detailed_users:
            users_list.append({
                'id': user.id,
                'email': user.email,
                'first_name': user.first_name or '',
                'last_name': user.last_name or '',
                'phone': getattr(user, 'phone', '') or '',
                'orders_count': user.orders_count or 0,
                'total_spent': user.total_spent or Decimal('0.00'),
                'date_joined': user.date_joined,
                'last_login': user.last_login,
                'is_active': user.is_active,
            })
        
        data = {
            'new_users': new_users,
            'new_users_change': calculate_change(new_users, prev_new_users),
            'total_users': total_users,
            'active_users': active_users,
            'users': users_list,
            'date_from': date_from,
            'date_to': date_to,
        }
        
        # Serialize the data
        serializer = UsersReportSerializer(data=data)
        if serializer.is_valid():
            return success_response(
                data=serializer.validated_data,
                message=_('تم جلب تقرير المستخدمين / Users report retrieved')
            )
        else:
            # If validation fails, return data directly
            return success_response(
                data=data,
                message=_('تم جلب تقرير المستخدمين / Users report retrieved')
            )


# =============================================================================
# Commissions Report View
# عرض تقرير العمولات
# =============================================================================

class CommissionsReportView(APIView):
    """
    Get commissions report.
    الحصول على تقرير العمولات.
    """
    
    permission_classes = [IsAdminUser]
    throttle_classes = [AdminUserRateThrottle]
    
    @extend_schema(
        summary='Commissions Report',
        description='Get commissions report with statistics',
        parameters=[
            OpenApiParameter(
                name='date_range',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Date range: 7days, 30days, 90days, year',
                enum=['7days', '30days', '90days', 'year'],
                default='30days'
            ),
        ],
        responses={
            200: CommissionsReportSerializer,
        },
        tags=['Admin Reports'],
    )
    def get(self, request):
        """
        Get commissions report.
        الحصول على تقرير العمولات.
        """
        from orders.models import Order
        
        date_range = request.query_params.get('date_range', '30days')
        date_from, date_to = get_date_range(date_range)
        date_from_dt = timezone.make_aware(datetime.combine(date_from, datetime.min.time()))
        date_to_dt = timezone.make_aware(datetime.combine(date_to, datetime.max.time()))
        
        # Get previous period
        prev_date_from, prev_date_to = get_previous_period(date_from, date_to)
        prev_date_from_dt = timezone.make_aware(datetime.combine(prev_date_from, datetime.min.time()))
        prev_date_to_dt = timezone.make_aware(datetime.combine(prev_date_to, datetime.max.time()))
        
        # Current period commissions
        current_orders = Order.objects.filter(
            created_at__gte=date_from_dt,
            created_at__lte=date_to_dt,
            status__in=['delivered', 'confirmed', 'shipped']
        ).select_related('user').prefetch_related('items__product_variant__product__vendor')
        
        total_commissions = current_orders.aggregate(total=Sum('platform_commission'))['total'] or Decimal('0.00')
        total_orders = current_orders.count()
        avg_commission_per_order = (total_commissions / total_orders) if total_orders > 0 else Decimal('0.00')
        
        # Previous period commissions
        previous_orders = Order.objects.filter(
            created_at__gte=prev_date_from_dt,
            created_at__lte=prev_date_to_dt,
            status__in=['delivered', 'confirmed', 'shipped']
        )
        
        prev_commissions = previous_orders.aggregate(total=Sum('platform_commission'))['total'] or Decimal('0.00')
        
        # Get detailed commissions list
        commissions_list = []
        for order in current_orders[:100]:  # Limit to 100 orders
            # Get vendor name from first order item
            vendor_name = 'غير معروف'
            if order.items.exists():
                first_item = order.items.first()
                if first_item and first_item.product_variant and first_item.product_variant.product:
                    vendor_name = first_item.product_variant.product.vendor.name if first_item.product_variant.product.vendor else 'غير معروف'
            
            commission_percentage = 10.0  # Default commission percentage
            if order.subtotal > 0:
                commission_percentage = float((order.platform_commission / order.subtotal) * 100)
            
            commissions_list.append({
                'order_id': order.id,
                'order_number': order.order_number or f"ORD-{order.id:06d}",
                'customer_name': order.customer_name or (order.user.get_full_name() if order.user else _('ضيف / Guest')),
                'vendor_name': vendor_name,
                'order_total': order.total,
                'commission_amount': order.platform_commission,
                'commission_percentage': round(commission_percentage, 2),
                'created_at': order.created_at,
            })
        
        data = {
            'total_commissions': total_commissions,
            'commissions_change': calculate_change(float(total_commissions), float(prev_commissions)),
            'total_orders': total_orders,
            'avg_commission_per_order': avg_commission_per_order,
            'commissions': commissions_list,
            'date_from': date_from,
            'date_to': date_to,
        }
        
        # Serialize the data
        serializer = CommissionsReportSerializer(data=data)
        if serializer.is_valid():
            return success_response(
                data=serializer.validated_data,
                message=_('تم جلب تقرير العمولات / Commissions report retrieved')
            )
        else:
            # If validation fails, return data directly
            return success_response(
                data=data,
                message=_('تم جلب تقرير العمولات / Commissions report retrieved')
            )


# =============================================================================
# Export Report View
# عرض تصدير التقرير
# =============================================================================

class ExportReportView(APIView):
    """
    Export report as Word document.
    تصدير التقرير كملف Word.
    """
    
    permission_classes = [IsAdminUser]
    throttle_classes = [AdminUserRateThrottle]
    
    @extend_schema(
        summary='Export Report',
        description='Export report as Word document (.docx)',
        parameters=[
            OpenApiParameter(
                name='type',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Report type: sales, products, users, commissions',
                enum=['sales', 'products', 'users', 'commissions'],
                required=True
            ),
            OpenApiParameter(
                name='date_range',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Date range: 7days, 30days, 90days, year',
                enum=['7days', '30days', '90days', 'year'],
                default='30days'
            ),
        ],
        responses={
            200: OpenApiResponse(
                description='Word document file',
                response={'application/vnd.openxmlformats-officedocument.wordprocessingml.document': bytes}
            ),
        },
        tags=['Admin Reports'],
    )
    def get(self, request):
        """
        Export report as Word.
        تصدير التقرير كملف Word.
        """
        report_type = request.query_params.get('type')
        date_range = request.query_params.get('date_range', '30days')
        
        # Get report data based on type
        if report_type == 'sales':
            # Get sales report data
            sales_view = SalesReportView()
            sales_view.request = request
            sales_response = sales_view.get(request)
            
            if not sales_response.data.get('success'):
                return error_response(
                    message=_('فشل في جلب بيانات التقرير / Failed to retrieve report data')
                )
            
            report_data = sales_response.data['data']
            doc = self._create_sales_word_document(report_data, date_range)
            
        elif report_type == 'products':
            products_view = ProductsReportView()
            products_view.request = request
            products_response = products_view.get(request)
            
            if not products_response.data.get('success'):
                return error_response(
                    message=_('فشل في جلب بيانات التقرير / Failed to retrieve report data')
                )
            
            report_data = products_response.data['data']
            doc = self._create_products_word_document(report_data, date_range)
            
        elif report_type == 'users':
            users_view = UsersReportView()
            users_view.request = request
            users_response = users_view.get(request)
            
            if not users_response.data.get('success'):
                return error_response(
                    message=_('فشل في جلب بيانات التقرير / Failed to retrieve report data')
                )
            
            report_data = users_response.data['data']
            doc = self._create_users_word_document(report_data, date_range)
            
        elif report_type == 'commissions':
            commissions_view = CommissionsReportView()
            commissions_view.request = request
            commissions_response = commissions_view.get(request)
            
            if not commissions_response.data.get('success'):
                return error_response(
                    message=_('فشل في جلب بيانات التقرير / Failed to retrieve report data')
                )
            
            report_data = commissions_response.data['data']
            doc = self._create_commissions_word_document(report_data, date_range)
            
        else:
            return error_response(
                message=_('نوع التقرير غير صحيح / Invalid report type')
            )
        
        # Save document to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        from django.http import HttpResponse
        
        filename = f'report_{report_type}_{date_range}_{datetime.now().strftime("%Y%m%d")}.docx'
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    def _create_sales_word_document(self, data, date_range):
        """Create Word document for sales report"""
        doc = Document()
        
        # Title
        title = doc.add_heading('تقرير المبيعات', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Date range
        date_para = doc.add_paragraph(f'الفترة: {data["date_from"]} إلى {data["date_to"]}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Summary
        doc.add_heading('ملخص الإحصائيات', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        summary_data = [
            ('إجمالي الإيرادات', f'{float(data["total_revenue"]):,.2f}'),
            ('إجمالي الطلبات', f'{int(data["total_orders"])}'),
            ('متوسط قيمة الطلب', f'{float(data["avg_order_value"]):,.2f}'),
            ('المستخدمين الجدد', f'{int(data["new_users"])}'),
        ]
        
        table = doc.add_table(rows=len(summary_data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(summary_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Daily sales
        if data.get('daily_sales'):
            doc.add_heading('المبيعات اليومية', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            daily_table = doc.add_table(rows=len(data['daily_sales']) + 1, cols=2)
            daily_table.style = 'Light Grid Accent 1'
            
            # Header
            daily_table.rows[0].cells[0].text = 'اليوم'
            daily_table.rows[0].cells[1].text = 'المبيعات'
            
            for i, daily in enumerate(data['daily_sales'], 1):
                daily_table.rows[i].cells[0].text = daily['day']
                daily_table.rows[i].cells[1].text = f'{float(daily["sales"]):,.2f}'
        
        # Detailed orders
        if data.get('orders') and len(data.get('orders', [])) > 0:
            doc.add_heading('قائمة الطلبات التفصيلية', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            orders_table = doc.add_table(rows=min(len(data['orders']), 50) + 1, cols=6)
            orders_table.style = 'Light Grid Accent 1'
            
            # Header
            orders_table.rows[0].cells[0].text = 'رقم الطلب'
            orders_table.rows[0].cells[1].text = 'العميل'
            orders_table.rows[0].cells[2].text = 'الحالة'
            orders_table.rows[0].cells[3].text = 'عدد العناصر'
            orders_table.rows[0].cells[4].text = 'الإجمالي'
            orders_table.rows[0].cells[5].text = 'التاريخ'
            
            for i, order in enumerate(data['orders'][:50], 1):  # Limit to 50 orders in Word
                orders_table.rows[i].cells[0].text = order.get('order_number', '')
                orders_table.rows[i].cells[1].text = order.get('customer_name', '')
                orders_table.rows[i].cells[2].text = order.get('status_display', order.get('status', ''))
                orders_table.rows[i].cells[3].text = str(order.get('items_count', 0))
                orders_table.rows[i].cells[4].text = f'{float(order.get("total", 0)):,.2f}'
                orders_table.rows[i].cells[5].text = str(order.get('created_at', ''))[:10]  # Date only
        
        return doc
    
    def _create_products_word_document(self, data, date_range):
        """Create Word document for products report"""
        doc = Document()
        
        title = doc.add_heading('تقرير المنتجات', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        date_para = doc.add_paragraph(f'الفترة: {data["date_from"]} إلى {data["date_to"]}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Top products
        if data.get('top_products'):
            doc.add_heading('المنتجات الأكثر مبيعاً', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            top_table = doc.add_table(rows=len(data['top_products']) + 1, cols=6)
            top_table.style = 'Light Grid Accent 1'
            
            top_table.rows[0].cells[0].text = 'المنتج'
            top_table.rows[0].cells[1].text = 'البائع'
            top_table.rows[0].cells[2].text = 'الفئة'
            top_table.rows[0].cells[3].text = 'المبيعات'
            top_table.rows[0].cells[4].text = 'الإيرادات'
            top_table.rows[0].cells[5].text = 'المخزون'
            
            for i, product in enumerate(data['top_products'], 1):
                top_table.rows[i].cells[0].text = product['name']
                top_table.rows[i].cells[1].text = product.get('vendor_name', 'غير معروف')
                top_table.rows[i].cells[2].text = product.get('category_name', 'أخرى')
                top_table.rows[i].cells[3].text = str(product['sales'])
                top_table.rows[i].cells[4].text = f'{float(product["revenue"]):,.2f}'
                top_table.rows[i].cells[5].text = str(product.get('stock_quantity', 0))
        
        # Sales by category
        if data.get('sales_by_category'):
            doc.add_heading('المبيعات حسب الفئة', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            cat_table = doc.add_table(rows=len(data['sales_by_category']) + 1, cols=3)
            cat_table.style = 'Light Grid Accent 1'
            
            cat_table.rows[0].cells[0].text = 'الفئة'
            cat_table.rows[0].cells[1].text = 'المبيعات'
            cat_table.rows[0].cells[2].text = 'النسبة %'
            
            for i, cat in enumerate(data['sales_by_category'], 1):
                cat_table.rows[i].cells[0].text = cat['category']
                cat_table.rows[i].cells[1].text = f'{float(cat["sales"]):,.2f}'
                cat_table.rows[i].cells[2].text = f'{float(cat["percentage"])}%'
        
        return doc
    
    def _create_users_word_document(self, data, date_range):
        """Create Word document for users report"""
        doc = Document()
        
        title = doc.add_heading('تقرير المستخدمين', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        date_para = doc.add_paragraph(f'الفترة: {data["date_from"]} إلى {data["date_to"]}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_heading('ملخص الإحصائيات', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        summary_data = [
            ('المستخدمين الجدد', f'{int(data["new_users"])}'),
            ('إجمالي المستخدمين', f'{int(data["total_users"])}'),
            ('المستخدمين النشطين', f'{int(data["active_users"])}'),
        ]
        
        table = doc.add_table(rows=len(summary_data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(summary_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        return doc
    
    def _create_commissions_word_document(self, data, date_range):
        """Create Word document for commissions report"""
        doc = Document()
        
        title = doc.add_heading('تقرير العمولات', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        date_para = doc.add_paragraph(f'الفترة: {data["date_from"]} إلى {data["date_to"]}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_heading('ملخص الإحصائيات', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        summary_data = [
            ('إجمالي العمولات', f'{float(data["total_commissions"]):,.2f}'),
            ('إجمالي الطلبات', f'{int(data["total_orders"])}'),
            ('متوسط العمولة لكل طلب', f'{float(data["avg_commission_per_order"]):,.2f}'),
        ]
        
        table = doc.add_table(rows=len(summary_data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(summary_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        return doc

