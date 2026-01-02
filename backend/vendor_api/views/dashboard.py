"""
Vendor Dashboard Views
عروض لوحة تحكم البائع

هذا الملف يحتوي على جميع views للـ Vendor Dashboard.
This file contains all Vendor Dashboard views.

Endpoints:
- GET /api/v1/vendor/dashboard/overview/        - KPIs and statistics
- GET /api/v1/vendor/dashboard/reports/export/ - Export report as Word
"""

from rest_framework.views import APIView
from django.utils import timezone
from django.utils.translation import gettext_lazy as _
from django.db.models import Sum, Count, Q
from django.db.models.functions import TruncDate, TruncMonth
from datetime import timedelta, datetime
from decimal import Decimal
from drf_spectacular.utils import extend_schema, OpenApiParameter, OpenApiResponse
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse

from vendor_api.permissions import IsVendorUser, IsVendorOwner
from vendor_api.throttling import VendorUserRateThrottle
from vendor_api.serializers.dashboard import (
    VendorDashboardOverviewSerializer,
    VendorSalesChartSerializer,
)
from core.utils import success_response, error_response
from users.models import VendorUser
import hashlib


# =============================================================================
# Vendor Dashboard Overview View
# عرض نظرة عامة على لوحة تحكم البائع
# =============================================================================

class VendorDashboardOverviewView(APIView):
    """
    Get vendor dashboard overview statistics (KPIs).
    الحصول على إحصائيات نظرة عامة على لوحة تحكم البائع.
    
    Returns all key performance indicators for the authenticated vendor including:
    - Sales statistics (total, today, change percentage)
    - Order statistics (total, today, pending, processing, change percentage)
    - Product statistics (total, active, low stock, out of stock)
    - Shop visits statistics (total, today, change percentage)
    
    Security:
    - Only authenticated vendors can access
    - Returns data only for the vendor associated with the authenticated user
    - Uses throttling to prevent abuse
    
    الأمان:
    - فقط البائعون المسجلون يمكنهم الوصول
    - يعيد البيانات فقط للبائع المرتبط بالمستخدم المسجل
    - يستخدم تحديد المعدل لمنع الإساءة
    """
    
    permission_classes = [IsVendorUser, IsVendorOwner]
    throttle_classes = [VendorUserRateThrottle]
    
    @extend_schema(
        summary='Vendor Dashboard Overview',
        description='Get all KPIs for the vendor dashboard',
        responses={
            200: VendorDashboardOverviewSerializer,
        },
        tags=['Vendor Dashboard'],
    )
    def get(self, request):
        """
        Get vendor dashboard overview statistics.
        الحصول على إحصائيات نظرة عامة على لوحة تحكم البائع.
        """
        # Import models here to avoid circular imports
        # استيراد النماذج هنا لتجنب الاستيراد الدائري
        from orders.models import Order, OrderItem
        from products.models import Product, ProductVariant
        
        # Get vendor associated with the authenticated user
        # الحصول على البائع المرتبط بالمستخدم المسجل
        try:
            vendor_user = VendorUser.objects.select_related('vendor').get(user=request.user)
            vendor = vendor_user.vendor
        except VendorUser.DoesNotExist:
            return success_response(
                data={},
                message=_('لا يوجد بائع مرتبط بهذا المستخدم / No vendor associated with this user'),
                status_code=404
            )
        
        # Time calculations
        # حسابات الوقت
        now = timezone.now()
        today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        week_start = today_start - timedelta(days=today_start.weekday())
        month_start = today_start.replace(day=1)
        last_month_start = (month_start - timedelta(days=1)).replace(day=1)
        
        # =================================================================
        # Revenue Statistics (Sales)
        # إحصائيات الإيرادات (المبيعات)
        # =================================================================
        
        # Get all orders and order items for this vendor (using denormalized vendor field)
        # الحصول على جميع الطلبات وعناصر الطلب لهذا البائع (باستخدام حقل vendor المطبيع)
        vendor_order_items = OrderItem.objects.filter(vendor=vendor)
        vendor_orders = Order.objects.filter(vendor=vendor)
        
        # Total sales (all time) - only completed/delivered orders
        # إجمالي المبيعات (كل الأوقات) - فقط الطلبات المكتملة/المسلمة
        # Calculate total sales correctly (sum of price * quantity for each item)
        # حساب إجمالي المبيعات بشكل صحيح (مجموع السعر * الكمية لكل عنصر)
        completed_items = vendor_order_items.filter(
            order__status__in=['delivered', 'completed']
        )
        total_sales = sum(
            Decimal(str(item.price)) * Decimal(str(item.quantity))
            for item in completed_items
        ) or Decimal('0.00')
        
        # This month's sales
        # مبيعات هذا الشهر
        this_month_items = vendor_order_items.filter(
            order__status__in=['delivered', 'completed'],
            order__created_at__gte=month_start
        )
        this_month_sales = sum(
            Decimal(str(item.price)) * Decimal(str(item.quantity))
            for item in this_month_items
        ) or Decimal('0.00')
        
        # Last month's sales
        # مبيعات الشهر الماضي
        last_month_items = vendor_order_items.filter(
            order__status__in=['delivered', 'completed'],
            order__created_at__gte=last_month_start,
            order__created_at__lt=month_start
        )
        last_month_sales = sum(
            Decimal(str(item.price)) * Decimal(str(item.quantity))
            for item in last_month_items
        ) or Decimal('0.00')
        
        # Sales change percentage
        # نسبة التغيير في المبيعات
        if last_month_sales > 0:
            sales_change = float(
                ((this_month_sales - last_month_sales) / last_month_sales) * 100
            )
        else:
            sales_change = 100.0 if this_month_sales > 0 else 0.0
        
        # Today's sales
        # مبيعات اليوم
        today_items = vendor_order_items.filter(
            order__created_at__gte=today_start
        )
        today_sales = sum(
            Decimal(str(item.price)) * Decimal(str(item.quantity))
            for item in today_items
        ) or Decimal('0.00')
        
        # =================================================================
        # Order Statistics
        # إحصائيات الطلبات
        # =================================================================
        
        total_orders = vendor_orders.count()
        today_orders = vendor_orders.filter(created_at__gte=today_start).count()
        pending_orders = vendor_orders.filter(status='pending').count()
        processing_orders = vendor_orders.filter(status__in=['processing', 'confirmed']).count()
        
        # This month's orders
        this_month_orders = vendor_orders.filter(created_at__gte=month_start).count()
        last_month_orders = vendor_orders.filter(
            created_at__gte=last_month_start,
            created_at__lt=month_start
        ).count()
        
        if last_month_orders > 0:
            orders_change = ((this_month_orders - last_month_orders) / last_month_orders) * 100
        else:
            orders_change = 100.0 if this_month_orders > 0 else 0.0
        
        # =================================================================
        # Product Statistics
        # إحصائيات المنتجات
        # =================================================================
        
        vendor_products = Product.objects.filter(vendor=vendor)
        total_products = vendor_products.count()
        active_products = vendor_products.filter(is_active=True).count()
        
        # Low stock products (stock < 10) - check through variants
        # منتجات بمخزون منخفض (مخزون < 10) - التحقق من خلال المتغيرات
        low_stock_products = 0
        out_of_stock_products = 0
        
        for product in vendor_products.filter(is_active=True):
            variants = ProductVariant.objects.filter(product=product)
            if not variants.exists():
                out_of_stock_products += 1
                continue
            
            # Check if any variant has stock
            # التحقق من وجود مخزون في أي متغير
            has_stock = False
            has_low_stock = False
            
            for variant in variants:
                # Assuming stock is stored in variant or calculated
                # افتراض أن المخزون مخزن في المتغير أو محسوب
                # TODO: Adjust based on actual stock field location
                stock = getattr(variant, 'stock', 0) if hasattr(variant, 'stock') else 0
                
                if stock > 0:
                    has_stock = True
                    if stock < 10:
                        has_low_stock = True
                        break
            
            if not has_stock:
                out_of_stock_products += 1
            elif has_low_stock:
                low_stock_products += 1
        
        # =================================================================
        # Shop Visits Statistics
        # إحصائيات زيارات المتجر
        # =================================================================
        
        # TODO: Implement shop visits tracking when analytics system is ready
        # TODO: تنفيذ تتبع زيارات المتجر عندما يكون نظام التحليلات جاهزاً
        
        # For now, return placeholder values
        # في الوقت الحالي، إرجاع قيم مؤقتة
        total_visits = 0
        today_visits = 0
        visits_change = 0.0
        
        # =================================================================
        # Response Rate Statistics
        # إحصائيات معدل الاستجابة
        # =================================================================
        
        # Calculate response rate based on completed/delivered orders
        # حساب معدل الاستجابة بناءً على الطلبات المكتملة/المسلمة
        # Response rate = (completed + delivered orders) / total orders * 100
        # معدل الاستجابة = (الطلبات المكتملة + المسلمة) / إجمالي الطلبات * 100
        
        # Strategy: Try to use last 30 days, if no data, use all orders
        # الاستراتيجية: محاولة استخدام آخر 30 يوم، إذا لم تكن هناك بيانات، استخدم جميع الطلبات
        
        # First, try last 30 days (more relevant for current performance)
        # أولاً، محاولة آخر 30 يوم (أكثر صلة بالأداء الحالي)
        thirty_days_ago = now - timedelta(days=30)
        recent_orders = vendor_orders.filter(created_at__gte=thirty_days_ago)
        total_recent_orders = recent_orders.count()
        
        # If we have recent orders, use them for calculation
        # إذا كانت لدينا طلبات حديثة، استخدمها للحساب
        if total_recent_orders > 0:
            # Count completed/delivered orders from recent period
            # عد الطلبات المكتملة/المسلمة من الفترة الحديثة
            completed_recent_orders = recent_orders.filter(
                status__in=['delivered', 'completed']
            ).count()
            
            # Calculate response rate percentage
            # حساب نسبة معدل الاستجابة
            response_rate = round((completed_recent_orders / total_recent_orders) * 100, 1)
            
            # Ensure response rate is between 0 and 100
            # التأكد من أن معدل الاستجابة بين 0 و 100
            response_rate = max(0.0, min(100.0, response_rate))
        else:
            # No recent orders, try to use all orders (fallback)
            # لا توجد طلبات حديثة، محاولة استخدام جميع الطلبات (بديل)
            total_all_orders = vendor_orders.count()
            
            if total_all_orders > 0:
                # Count completed/delivered orders from all time
                # عد الطلبات المكتملة/المسلمة من كل الأوقات
                completed_all_orders = vendor_orders.filter(
                    status__in=['delivered', 'completed']
                ).count()
                
                # Calculate response rate from all orders
                # حساب معدل الاستجابة من جميع الطلبات
                response_rate = round((completed_all_orders / total_all_orders) * 100, 1)
                
                # Ensure response rate is between 0 and 100
                # التأكد من أن معدل الاستجابة بين 0 و 100
                response_rate = max(0.0, min(100.0, response_rate))
            else:
                # No orders at all - return None (no data available)
                # لا توجد طلبات على الإطلاق - إرجاع None (لا توجد بيانات متاحة)
                response_rate = None
        
        # =================================================================
        # Build Response
        # بناء الاستجابة
        # =================================================================
        
        data = {
            # Sales (Revenue)
            'total_sales': str(total_sales),  # Convert Decimal to string for serializer
            'total_sales_change': round(sales_change, 1),
            'today_sales': str(today_sales),  # Convert Decimal to string for serializer
            
            # Orders
            'total_orders': total_orders,
            'total_orders_change': round(orders_change, 1),
            'today_orders': today_orders,
            'pending_orders': pending_orders,
            'processing_orders': processing_orders,
            
            # Products
            'total_products': total_products,
            'active_products': active_products,
            'low_stock_products': low_stock_products,
            'out_of_stock_products': out_of_stock_products,
            
            # Shop Visits
            'total_visits': total_visits,
            'total_visits_change': round(visits_change, 1),
            'today_visits': today_visits,
            
            # Response Rate
            'response_rate': response_rate,
        }
        
        # Validate and serialize the data
        # التحقق من البيانات وتسلسلها
        serializer = VendorDashboardOverviewSerializer(data=data)
        
        if serializer.is_valid(raise_exception=True):
            return success_response(
                data=serializer.validated_data,
                message=_('تم جلب إحصائيات لوحة التحكم بنجاح / Dashboard statistics retrieved successfully')
            )


# =============================================================================
# Vendor Sales Chart View
# عرض رسم بياني مبيعات البائع
# =============================================================================

class VendorSalesChartView(APIView):
    """
    Get sales chart data for vendor dashboard.
    الحصول على بيانات رسم بياني المبيعات للوحة تحكم البائع.
    
    Returns sales data grouped by time period (day/week/month).
    Only includes sales from this vendor's products.
    
    يعيد بيانات المبيعات مجمعة حسب الفترة الزمنية (يوم/أسبوع/شهر).
    يتضمن فقط المبيعات من منتجات هذا البائع.
    
    Security:
    - Only authenticated vendors can access
    - Returns data only for the vendor's products
    - Uses throttling to prevent abuse
    
    الأمان:
    - فقط البائعون المسجلون يمكنهم الوصول
    - يعيد البيانات فقط لمنتجات البائع
    - يستخدم تحديد المعدل لمنع الإساءة
    """
    
    permission_classes = [IsVendorUser, IsVendorOwner]
    throttle_classes = [VendorUserRateThrottle]
    
    @extend_schema(
        summary='Vendor Sales Chart Data',
        description='Get sales chart data for vendor dashboard with different time periods',
        parameters=[
            OpenApiParameter(
                name='period',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Time period: week, month, or year',
                enum=['week', 'month', 'year'],
                default='month'
            ),
        ],
        tags=['Vendor Dashboard'],
    )
    def get(self, request):
        """
        Get sales chart data.
        الحصول على بيانات رسم بياني المبيعات.
        """
        from orders.models import Order, OrderItem
        
        # Get vendor associated with the authenticated user
        # الحصول على البائع المرتبط بالمستخدم المسجل
        try:
            vendor_user = VendorUser.objects.select_related('vendor').get(user=request.user)
            vendor = vendor_user.vendor
        except VendorUser.DoesNotExist:
            return error_response(
                message=_('لا يوجد بائع مرتبط بهذا المستخدم / No vendor associated with this user')
            )
        
        # Get period parameter
        # الحصول على معامل الفترة
        period = request.query_params.get('period', 'month')
        if period not in ['week', 'month', 'year']:
            period = 'month'  # Default to month
        
        now = timezone.now()
        
        # Determine date range and grouping based on period
        # تحديد نطاق التاريخ والتجميع بناءً على الفترة
        if period == 'week':
            # Last 7 days, grouped by day
            # آخر 7 أيام، مجمعة حسب اليوم
            start_date = now - timedelta(days=7)
            trunc_func = TruncDate
            date_format = '%Y-%m-%d'
            label_format = '%d/%m'  # DD/MM for display
        elif period == 'month':
            # Last 30 days, grouped by day
            # آخر 30 يوم، مجمعة حسب اليوم
            start_date = now - timedelta(days=30)
            trunc_func = TruncDate
            date_format = '%Y-%m-%d'
            label_format = '%d/%m'  # DD/MM for display
        else:  # year
            # Last 12 months, grouped by month
            # آخر 12 شهر، مجمعة حسب الشهر
            start_date = now - timedelta(days=365)
            trunc_func = TruncMonth
            date_format = '%Y-%m'
            label_format = '%b %Y'  # Month Year for display
        
        # Get all order items that belong to this vendor (using denormalized vendor field)
        # الحصول على جميع عناصر الطلب التي تنتمي لهذا البائع (باستخدام حقل vendor المطبيع)
        vendor_order_items = OrderItem.objects.filter(
            vendor=vendor,
            order__created_at__gte=start_date,
            order__status__in=['delivered', 'completed', 'processing', 'pending', 'confirmed']
        ).select_related('order', 'product_variant', 'product_variant__product')
        
        # Group by date and calculate totals
        # التجميع حسب التاريخ وحساب الإجماليات
        # We need to group by order date, not item date
        # نحتاج للتجميع حسب تاريخ الطلب، وليس تاريخ العنصر
        
        # Get unique orders with their dates (optimized query)
        # الحصول على الطلبات الفريدة مع تواريخها (استعلام محسّن)
        order_ids = vendor_order_items.values_list('order_id', flat=True).distinct()
        
        if not order_ids:
            # No orders, return empty data
            # لا توجد طلبات، إرجاع بيانات فارغة
            labels = []
            revenue = []
            orders_count = []
            
            # Fill with empty data for the period
            # ملء ببيانات فارغة للفترة
            if period == 'week':
                days = 7
                for i in range(days):
                    date = (now - timedelta(days=days - i - 1)).replace(hour=0, minute=0, second=0, microsecond=0)
                    labels.append(date.strftime(label_format))
                    revenue.append('0')
                    orders_count.append(0)
            elif period == 'month':
                days = 30
                for i in range(days):
                    date = (now - timedelta(days=days - i - 1)).replace(hour=0, minute=0, second=0, microsecond=0)
                    labels.append(date.strftime(label_format))
                    revenue.append('0')
                    orders_count.append(0)
            else:  # year
                months = 12
                for i in range(months):
                    month_date = (now - timedelta(days=30 * (months - i - 1))).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                    labels.append(month_date.strftime(label_format))
                    revenue.append('0')
                    orders_count.append(0)
            
            data = {
                'labels': labels,
                'revenue': revenue,
                'orders': orders_count,
                'period': period,
            }
            
            return success_response(
                data=data,
                message=_('تم جلب بيانات الرسم البياني / Chart data retrieved')
            )
        
        # Get orders with their dates (optimized)
        # الحصول على الطلبات مع تواريخها (محسّن)
        orders = Order.objects.filter(
            id__in=order_ids
        ).annotate(
            date=trunc_func('created_at')
        ).values('id', 'date')
        
        # Pre-calculate revenue per order item (more efficient)
        # حساب الإيرادات مسبقاً لكل عنصر طلب (أكثر كفاءة)
        items_revenue = {}
        for item in vendor_order_items:
            order_id = item.order_id
            item_revenue = Decimal(str(item.price)) * Decimal(str(item.quantity))
            
            if order_id not in items_revenue:
                items_revenue[order_id] = Decimal('0.00')
            items_revenue[order_id] += item_revenue
        
        # Build a map of date -> aggregated data
        # بناء خريطة من التاريخ -> البيانات المجمعة
        date_items_map = {}
        for order_data in orders:
            order_date = order_data['date']
            order_id = order_data['id']
            
            # Calculate total revenue for this date
            # حساب إجمالي الإيرادات لهذا التاريخ
            if order_date not in date_items_map:
                date_items_map[order_date] = {
                    'revenue': Decimal('0.00'),
                    'orders_count': 0,
                    'order_ids': set()
                }
            
            # Add revenue from this order
            # إضافة الإيرادات من هذا الطلب
            order_revenue = items_revenue.get(order_id, Decimal('0.00'))
            date_items_map[order_date]['revenue'] += order_revenue
            
            # Count unique orders per date
            # عد الطلبات الفريدة لكل تاريخ
            if order_id not in date_items_map[order_date]['order_ids']:
                date_items_map[order_date]['order_ids'].add(order_id)
                date_items_map[order_date]['orders_count'] += 1
        
        # Build labels and data arrays
        # بناء مصفوفات التسميات والبيانات
        labels = []
        revenue = []
        orders_count = []
        
        # Sort by date
        # الترتيب حسب التاريخ
        sorted_dates = sorted(date_items_map.keys())
        
        for date in sorted_dates:
            # Format label for display
            # تنسيق التسمية للعرض
            if period == 'year':
                # For year, use month format
                # للسنة، استخدم تنسيق الشهر
                label = date.strftime(label_format)
            else:
                # For week/month, use day format
                # للأسبوع/الشهر، استخدم تنسيق اليوم
                label = date.strftime(label_format)
            
            labels.append(label)
            revenue.append(str(date_items_map[date]['revenue']))
            orders_count.append(date_items_map[date]['orders_count'])
        
        # If no data, return empty arrays
        # إذا لم تكن هناك بيانات، إرجاع مصفوفات فارغة
        if not labels:
            # Fill with empty data for the period
            # ملء ببيانات فارغة للفترة
            if period == 'week':
                days = 7
                for i in range(days):
                    date = (now - timedelta(days=days - i - 1)).replace(hour=0, minute=0, second=0, microsecond=0)
                    labels.append(date.strftime(label_format))
                    revenue.append('0')
                    orders_count.append(0)
            elif period == 'month':
                days = 30
                for i in range(days):
                    date = (now - timedelta(days=days - i - 1)).replace(hour=0, minute=0, second=0, microsecond=0)
                    labels.append(date.strftime(label_format))
                    revenue.append('0')
                    orders_count.append(0)
            else:  # year
                months = 12
                for i in range(months):
                    # Calculate month date
                    month_date = (now - timedelta(days=30 * (months - i - 1))).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
                    labels.append(month_date.strftime(label_format))
                    revenue.append('0')
                    orders_count.append(0)
        
        # Build response data
        # بناء بيانات الاستجابة
        data = {
            'labels': labels,
            'revenue': revenue,
            'orders': orders_count,
            'period': period,
        }
        
        # Return data directly (already validated and formatted)
        # إرجاع البيانات مباشرة (محققة ومنسقة بالفعل)
        return success_response(
            data=data,
            message=_('تم جلب بيانات الرسم البياني / Chart data retrieved')
            )


# =============================================================================
# Vendor Dashboard Tips View
# عرض نصائح لوحة تحكم البائع
# =============================================================================

class VendorDashboardTipsView(APIView):
    """
    Get dynamic tips for vendor dashboard.
    الحصول على نصائح ديناميكية للوحة تحكم البائع.
    
    Returns actionable tips based on vendor's current status:
    - Low stock products
    - Out of stock products
    - Inactive products
    - Other relevant suggestions
    
    يعيد نصائح قابلة للتنفيذ بناءً على حالة البائع الحالية:
    - منتجات قليلة المخزون
    - منتجات نفد المخزون
    - منتجات غير نشطة
    - اقتراحات أخرى ذات صلة
    
    Security:
    - Only authenticated vendors can access
    - Returns tips only for the vendor's products
    - Uses throttling to prevent abuse
    
    الأمان:
    - فقط البائعون المسجلون يمكنهم الوصول
    - يعيد النصائح فقط لمنتجات البائع
    - يستخدم تحديد المعدل لمنع الإساءة
    """
    
    permission_classes = [IsVendorUser, IsVendorOwner]
    throttle_classes = [VendorUserRateThrottle]
    
    @extend_schema(
        summary='Vendor Dashboard Tips',
        description='Get dynamic tips for vendor dashboard based on current status',
        tags=['Vendor Dashboard'],
    )
    def get(self, request):
        """
        Get dashboard tips.
        الحصول على نصائح لوحة التحكم.
        """
        from products.models import Product, ProductVariant
        
        # Get vendor associated with the authenticated user
        # الحصول على البائع المرتبط بالمستخدم المسجل
        try:
            vendor_user = VendorUser.objects.select_related('vendor').get(user=request.user)
            vendor = vendor_user.vendor
        except VendorUser.DoesNotExist:
            return error_response(
                message=_('لا يوجد بائع مرتبط بهذا المستخدم / No vendor associated with this user')
            )
        
        # Get out of stock products (all variants have stock_quantity = 0)
        # الحصول على المنتجات التي نفد مخزونها (جميع المتغيرات لديها stock_quantity = 0)
        out_of_stock_products = Product.objects.filter(
            vendor=vendor,
            is_active=True
        ).annotate(
            total_stock=Sum('variants__stock_quantity')
        ).filter(
            total_stock=0
        ).prefetch_related('variants')[:5]  # Limit to 5 products
        
        # Get low stock products (has variants with stock_quantity < 10 and > 0)
        # الحصول على المنتجات قليلة المخزون (لديها متغيرات بمخزون < 10 و > 0)
        low_stock_products = Product.objects.filter(
            vendor=vendor,
            is_active=True,
            variants__stock_quantity__lt=10,
            variants__stock_quantity__gt=0
        ).distinct().prefetch_related('variants')[:5]  # Limit to 5 products
        
        # Get inactive products
        # الحصول على المنتجات غير النشطة
        inactive_products = Product.objects.filter(
            vendor=vendor,
            is_active=False
        )[:5]  # Limit to 5 products
        
        # Build tips based on priority
        # بناء النصائح بناءً على الأولوية
        tips = []
        
        # Priority 1: Out of stock products
        # الأولوية 1: المنتجات التي نفد مخزونها
        if out_of_stock_products.exists():
            product = out_of_stock_products.first()
            tip = {
                'type': 'out_of_stock',
                'priority': 1,
                'title_ar': 'منتج نفد مخزونه',
                'title_en': 'Out of Stock Product',
                'message_ar': f'لقد نفذت كمية <b>"{product.name}"</b>. قم بتحديث المخزون الآن لضمان عدم ضياع أي طلبات محتملة.',
                'message_en': f'The stock for <b>"{product.name}"</b> is out. Update inventory now to avoid losing potential orders.',
                'action_text_ar': 'تحديث المخزون',
                'action_text_en': 'Update Stock',
                'action_url': f'/vendor/products/{product.id}',
                'product_id': product.id,
                'product_name': product.name,
            }
            tips.append(tip)
        
        # Priority 2: Low stock products
        # الأولوية 2: المنتجات قليلة المخزون
        elif low_stock_products.exists():
            product = low_stock_products.first()
            # Get the variant with lowest stock
            # الحصول على المتغير بأقل مخزون
            variant = product.variants.filter(
                stock_quantity__lt=10,
                stock_quantity__gt=0
            ).order_by('stock_quantity').first()
            
            variant_name = f"{variant.color}"
            if variant.size:
                variant_name += f" - {variant.size}"
            
            tip = {
                'type': 'low_stock',
                'priority': 2,
                'title_ar': 'منتج قليل المخزون',
                'title_en': 'Low Stock Product',
                'message_ar': f'المخزون لـ <b>"{product.name} - {variant_name}"</b> منخفض ({variant.stock_quantity} قطع متبقية). قم بتحديث المخزون قريباً.',
                'message_en': f'Stock for <b>"{product.name} - {variant_name}"</b> is low ({variant.stock_quantity} items remaining). Update inventory soon.',
                'action_text_ar': 'تحديث المخزون',
                'action_text_en': 'Update Stock',
                'action_url': f'/vendor/products/{product.id}',
                'product_id': product.id,
                'product_name': product.name,
            }
            tips.append(tip)
        
        # Priority 3: Inactive products
        # الأولوية 3: المنتجات غير النشطة
        elif inactive_products.exists():
            product = inactive_products.first()
            tip = {
                'type': 'inactive',
                'priority': 3,
                'title_ar': 'منتج غير نشط',
                'title_en': 'Inactive Product',
                'message_ar': f'المنتج <b>"{product.name}"</b> غير نشط حالياً. قم بتفعيله لزيادة المبيعات.',
                'message_en': f'Product <b>"{product.name}"</b> is currently inactive. Activate it to increase sales.',
                'action_text_ar': 'تفعيل المنتج',
                'action_text_en': 'Activate Product',
                'action_url': f'/vendor/products/{product.id}',
                'product_id': product.id,
                'product_name': product.name,
            }
            tips.append(tip)
        
        # Default tip if no issues found
        # نصيحة افتراضية إذا لم توجد مشاكل
        if not tips:
            tip = {
                'type': 'general',
                'priority': 4,
                'title_ar': 'نصيحة عامة',
                'title_en': 'General Tip',
                'message_ar': 'أداؤك ممتاز! استمر في تحديث المخزون بانتظام للحفاظ على المبيعات.',
                'message_en': 'Your performance is excellent! Keep updating inventory regularly to maintain sales.',
                'action_text_ar': 'عرض المنتجات',
                'action_text_en': 'View Products',
                'action_url': '/vendor/products',
                'product_id': None,
                'product_name': None,
            }
            tips.append(tip)
        
        # Return the highest priority tip
        # إرجاع النصيحة ذات الأولوية الأعلى
        tip = tips[0]  # Already sorted by priority
        
        # Build response data
        # بناء بيانات الاستجابة
        data = {
            'type': tip['type'],
            'priority': tip['priority'],
            'title_ar': tip['title_ar'],
            'title_en': tip['title_en'],
            'message_ar': tip['message_ar'],
            'message_en': tip['message_en'],
            'action_text_ar': tip['action_text_ar'],
            'action_text_en': tip['action_text_en'],
            'action_url': tip['action_url'],
            'product_id': tip['product_id'],
            'product_name': tip['product_name'],
        }
        
        # Return data directly (already validated and formatted)
        # إرجاع البيانات مباشرة (محققة ومنسقة بالفعل)
        return success_response(
            data=data,
            message=_('تم جلب النصيحة / Tip retrieved')
        )


# =============================================================================
# Vendor Recent Orders View
# عرض الطلبات الأخيرة للبائع
# =============================================================================

class VendorRecentOrdersView(APIView):
    """
    Get recent orders for vendor dashboard.
    الحصول على الطلبات الأخيرة للوحة تحكم البائع.
    
    Returns the most recent orders that contain products from this vendor.
    يعيد آخر الطلبات التي تحتوي على منتجات من هذا البائع.
    
    Security:
    - Only authenticated vendors can access
    - Returns orders only for the vendor's products
    
    الأمان:
    - فقط البائعون المسجلون يمكنهم الوصول
    - يعيد الطلبات فقط لمنتجات البائع
    """
    
    permission_classes = [IsVendorUser, IsVendorOwner]
    throttle_classes = [VendorUserRateThrottle]
    
    @extend_schema(
        summary='Vendor Recent Orders',
        description='Get the most recent orders for this vendor',
        parameters=[
            OpenApiParameter(
                name='limit',
                type=int,
                location=OpenApiParameter.QUERY,
                description='Number of orders to return (default: 10, max: 50)',
                default=10
            ),
        ],
        tags=['Vendor Dashboard'],
    )
    def get(self, request):
        """
        Get recent orders.
        الحصول على الطلبات الأخيرة.
        """
        from orders.models import Order, OrderItem
        
        # Get vendor associated with the authenticated user
        # الحصول على البائع المرتبط بالمستخدم المسجل
        try:
            vendor_user = VendorUser.objects.select_related('vendor').get(user=request.user)
            vendor = vendor_user.vendor
        except VendorUser.DoesNotExist:
            return error_response(
                message=_('لا يوجد بائع مرتبط بهذا المستخدم / No vendor associated with this user')
            )
        
        limit = int(request.query_params.get('limit', 10))
        limit = min(limit, 50)  # Max 50 orders
        
        # Get all order items that belong to this vendor (using denormalized vendor field)
        # الحصول على جميع عناصر الطلب التي تنتمي لهذا البائع (باستخدام حقل vendor المطبيع)
        vendor_order_items = OrderItem.objects.filter(
            vendor=vendor
        ).select_related('order', 'order__user')
        
        # Get unique orders
        # الحصول على الطلبات الفريدة
        order_ids = vendor_order_items.values_list('order_id', flat=True).distinct()
        orders = (
            Order.objects.filter(id__in=order_ids)
            .select_related('user')
            .prefetch_related('items')
            .order_by('-created_at')[:limit]
        )
        
        # Status display mapping
        # ربط حالة العرض
        status_display_map = {
            'pending': _('قيد الانتظار / Pending'),
            'confirmed': _('مؤكد / Confirmed'),
            'shipped': _('تم الشحن / Shipped'),
            'delivered': _('تم التسليم / Delivered'),
            'cancelled': _('ملغي / Cancelled'),
        }
        
        # Helper function to generate customer_key
        # دالة مساعدة لإنشاء customer_key
        def generate_customer_key(vendor_id: int, user_id: int = None, guest_identifier: str = None) -> str:
            if user_id:
                raw_key = f"vendor_{vendor_id}_user_{user_id}"
            elif guest_identifier:
                raw_key = f"vendor_{vendor_id}_guest_{guest_identifier}"
            else:
                raw_key = f"vendor_{vendor_id}_unknown"
            hash_obj = hashlib.sha256(raw_key.encode())
            return hash_obj.hexdigest()[:16]
        
        # Build response data
        # بناء بيانات الاستجابة
        data = []
        
        for order in orders:
            # Get customer name
            # الحصول على اسم العميل
            if order.user:
                customer_name = f"{order.user.first_name} {order.user.last_name}".strip()
                if not customer_name:
                    customer_name = order.user.email.split('@')[0] if order.user.email else _('مستخدم / User')
                # Calculate customer_key for authenticated customer
                # حساب customer_key للعميل المسجل
                customer_key = generate_customer_key(vendor.id, order.user.id)
            else:
                customer_name = order.customer_name or _('ضيف / Guest')
                # Calculate customer_key for guest customer
                # حساب customer_key للعميل الضيف
                guest_identifier = order.customer_name or order.customer_phone or f"guest_{order.id}"
                customer_key = generate_customer_key(vendor.id, None, guest_identifier)
            
            # Calculate total for this vendor's items only
            # حساب الإجمالي لعناصر هذا البائع فقط
            vendor_items = vendor_order_items.filter(order=order)
            vendor_total = sum(
                Decimal(str(item.price)) * Decimal(str(item.quantity))
                for item in vendor_items
            ) or Decimal('0.00')
            
            data.append({
                'id': order.id,
                'order_number': order.order_number or f"ORD-{order.id:06d}",
                'customer_name': customer_name,
                'customer_key': customer_key,
                'total': str(vendor_total),
                'status': order.status,
                'status_display': status_display_map.get(order.status, order.status),
                'created_at': order.created_at,
            })
        
        # Return data directly (already validated and formatted)
        # إرجاع البيانات مباشرة (محققة ومنسقة بالفعل)
        return success_response(
            data=data,
            message=_('تم جلب الطلبات الأخيرة / Recent orders retrieved')
        )


# =============================================================================
# Vendor Report Export View
# عرض تصدير تقرير البائع
# =============================================================================

class VendorReportExportView(APIView):
    """
    Export vendor report as Word document.
    تصدير تقرير البائع كملف Word.
    
    This view exports the vendor dashboard statistics as a Word document.
    يعرض هذا العرض إحصائيات لوحة تحكم البائع كمستند Word.
    
    Security:
    - Only authenticated vendors can access
    - Returns data only for the vendor associated with the authenticated user
    
    الأمان:
    - فقط البائعون المسجلون يمكنهم الوصول
    - يعيد البيانات فقط للبائع المرتبط بالمستخدم المسجل
    """
    
    permission_classes = [IsVendorUser, IsVendorOwner]
    throttle_classes = [VendorUserRateThrottle]
    
    @extend_schema(
        summary='Export Vendor Report',
        description='Export vendor dashboard report as Word document (.docx)',
        parameters=[
            OpenApiParameter(
                name='date_range',
                type=str,
                location=OpenApiParameter.QUERY,
                description='Date range: 7days, 30days, 90days, year',
                enum=['7days', '30days', '90days', 'year'],
                default='30days'
            ),
        ],
        responses={
            200: OpenApiResponse(
                description='Word document file',
                response={'application/vnd.openxmlformats-officedocument.wordprocessingml.document': bytes}
            ),
        },
        tags=['Vendor Dashboard'],
    )
    def get(self, request):
        """
        Export vendor report as Word.
        تصدير تقرير البائع كملف Word.
        """
        date_range = request.query_params.get('date_range', '30days')
        
        # Get vendor associated with the authenticated user
        # الحصول على البائع المرتبط بالمستخدم المسجل
        try:
            vendor_user = VendorUser.objects.select_related('vendor').get(user=request.user)
            vendor = vendor_user.vendor
        except VendorUser.DoesNotExist:
            return error_response(
                message=_('لا يوجد بائع مرتبط بهذا المستخدم / No vendor associated with this user')
            )
        
        # Get dashboard overview data
        # الحصول على بيانات نظرة عامة على لوحة التحكم
        overview_view = VendorDashboardOverviewView()
        overview_view.request = request
        overview_response = overview_view.get(request)
        
        if not overview_response.data.get('success'):
            return error_response(
                message=_('فشل في جلب بيانات التقرير / Failed to retrieve report data')
            )
        
        dashboard_data = overview_response.data['data']
        
        # Create Word document
        # إنشاء مستند Word
        doc = self._create_vendor_word_document(dashboard_data, vendor, date_range)
        
        # Save document to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Create response
        filename = f'vendor_report_{date_range}_{datetime.now().strftime("%Y%m%d")}.docx'
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    def _create_vendor_word_document(self, data, vendor, date_range):
        """Create Word document for vendor report"""
        doc = Document()
        
        # Title
        title = doc.add_heading('تقرير البائع', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Vendor name
        vendor_para = doc.add_paragraph(f'اسم المتجر: {vendor.name}')
        vendor_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Date
        date_para = doc.add_paragraph(f'تاريخ التقرير: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Period
        period_labels = {
            '7days': 'آخر 7 أيام',
            '30days': 'آخر 30 يوم',
            '90days': 'آخر 90 يوم',
            'year': 'آخر سنة'
        }
        period_para = doc.add_paragraph(f'الفترة: {period_labels.get(date_range, date_range)}')
        period_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Empty line
        
        # Sales Summary
        doc.add_heading('ملخص المبيعات', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        sales_data = [
            ('إجمالي المبيعات', f'{float(data.get("total_sales", 0)):,.2f} ل.س'),
            ('التغيير من الشهر الماضي', f'{data.get("total_sales_change", 0):.1f}%'),
            ('مبيعات اليوم', f'{float(data.get("today_sales", 0)):,.2f} ل.س'),
        ]
        
        sales_table = doc.add_table(rows=len(sales_data), cols=2)
        sales_table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(sales_data):
            sales_table.rows[i].cells[0].text = label
            sales_table.rows[i].cells[1].text = str(value)
            sales_table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sales_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Empty line
        
        # Orders Summary
        doc.add_heading('ملخص الطلبات', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        orders_data = [
            ('إجمالي الطلبات', str(data.get('total_orders', 0))),
            ('التغيير من الشهر الماضي', f'{data.get("total_orders_change", 0):.1f}%'),
            ('طلبات اليوم', str(data.get('today_orders', 0))),
            ('طلبات قيد الانتظار', str(data.get('pending_orders', 0))),
            ('طلبات قيد المعالجة', str(data.get('processing_orders', 0))),
        ]
        
        orders_table = doc.add_table(rows=len(orders_data), cols=2)
        orders_table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(orders_data):
            orders_table.rows[i].cells[0].text = label
            orders_table.rows[i].cells[1].text = str(value)
            orders_table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            orders_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Empty line
        
        # Products Summary
        doc.add_heading('ملخص المنتجات', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        products_data = [
            ('إجمالي المنتجات', str(data.get('total_products', 0))),
            ('المنتجات النشطة', str(data.get('active_products', 0))),
            ('منتجات قليلة المخزون', str(data.get('low_stock_products', 0))),
            ('منتجات نفد المخزون', str(data.get('out_of_stock_products', 0))),
        ]
        
        products_table = doc.add_table(rows=len(products_data), cols=2)
        products_table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(products_data):
            products_table.rows[i].cells[0].text = label
            products_table.rows[i].cells[1].text = str(value)
            products_table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            products_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Empty line
        
        # Visits Summary
        doc.add_heading('ملخص الزيارات', 1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        visits_data = [
            ('إجمالي الزيارات', str(data.get('total_visits', 0))),
            ('التغيير من الأسبوع الماضي', f'{data.get("total_visits_change", 0):.1f}%'),
            ('زيارات اليوم', str(data.get('today_visits', 0))),
        ]
        
        visits_table = doc.add_table(rows=len(visits_data), cols=2)
        visits_table.style = 'Light Grid Accent 1'
        
        for i, (label, value) in enumerate(visits_data):
            visits_table.rows[i].cells[0].text = label
            visits_table.rows[i].cells[1].text = str(value)
            visits_table.rows[i].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            visits_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Footer
        doc.add_paragraph()  # Empty line
        footer = doc.add_paragraph('تم إنشاء هذا التقرير تلقائياً من نظام إدارة المتجر')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc

